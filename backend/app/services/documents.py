"""Document generation utilities.

write_docx_rich  — python-docx based writer with headings and bullets (ATS-friendly)
write_docx       — lightweight raw XML fallback (no extra dependency)
write_pdf        — minimal PDF-1.4 writer
"""

from __future__ import annotations

import html
import logging
import os
import shutil
import subprocess
import tempfile
import textwrap
import zipfile
from pathlib import Path

logger = logging.getLogger(__name__)


def latex_compiler_available() -> bool:
    return any(shutil.which(engine) for engine in ["latexmk", "xelatex", "pdflatex", "lualatex"])


def docx_pdf_converter_available() -> bool:
    return any(shutil.which(engine) for engine in ["soffice", "libreoffice"]) or _microsoft_word_available()


def _microsoft_word_available() -> bool:
    enabled = os.getenv("SEEKAPPLY_ENABLE_MS_WORD_PDF", "").strip().lower() in {"1", "true", "yes"}
    return enabled and Path("/Applications/Microsoft Word.app").exists() and shutil.which("osascript") is not None


# ---------------------------------------------------------------------------
# Rich DOCX writer (python-docx)
# ---------------------------------------------------------------------------

def write_docx_rich(path: Path, title: str, paragraphs: list[str]) -> None:
    """Write a formatted DOCX using python-docx.

    Paragraphs prefixed with ``__HEADING__`` are rendered as bold section
    headers. Lines starting with ``•`` or ``-`` are rendered as bullet items.
    Falls back to write_docx() if python-docx is not installed.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # --- Document-level style tweaks ---
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()  # spacer

        for line in paragraphs:
            if line.startswith("__HEADING__"):
                text = line[len("__HEADING__"):]
                heading_para = doc.add_paragraph()
                run = heading_para.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)  # dark blue
                # Add a thin rule effect via bottom border (simple approach)
            elif line.startswith("•") or line.startswith("-"):
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.add_run(line.lstrip("•-").strip())
            else:
                doc.add_paragraph(line)

        doc.save(path)

    except ImportError:
        # Fallback to raw XML writer
        write_docx(path, title, paragraphs)


def write_docx_from_template(
    path: Path,
    template_path: Path,
    *,
    job_title: str,
    company: str,
    paragraphs: list[str],
    emphasized_skills: list[str],
    selected_projects: list[dict],
) -> bool:
    """Create a tailored DOCX by making small edits to an uploaded Word resume.

    The goal is to keep the user's Word layout intact while replacing only the
    profile/summary and adding compact targeted skill evidence. The Projects
    section is preserved exactly so generated or inferred project text cannot
    look like unsupported resume evidence. If python-docx cannot load the
    template, the caller can fall back to the clean DOCX writer.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
    except ImportError:
        return False

    if not template_path.exists() or template_path.suffix.lower() != ".docx":
        return False

    def section_lines(header: str) -> list[str]:
        capture = False
        lines: list[str] = []
        for item in paragraphs:
            text = item.replace("__HEADING__", "").strip()
            if not text:
                continue
            if item.startswith("__HEADING__"):
                if capture:
                    break
                capture = text.lower() == header.lower()
                continue
            if capture:
                lines.append(text)
        return lines

    def all_paragraphs(document) -> list:
        items = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    items.extend(cell.paragraphs)
        return items

    def norm(value: str) -> str:
        return " ".join(value.lower().strip().split())

    def replace_text(paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def insert_after(paragraph, text: str) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        try:
            new_para.style = paragraph.style
        except Exception:
            pass
        new_para.add_run(text)
        return new_para

    def remove_paragraph(paragraph) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
            paragraph._p = paragraph._element = None

    def clean_previous_targeted_edits(document) -> None:
        paragraphs = all_paragraphs(document)
        remove_next_project_lines = 0
        for paragraph in list(paragraphs):
            text = norm(paragraph.text)
            if not text:
                continue
            if text.startswith("targeted focus for "):
                remove_paragraph(paragraph)
                continue
            if text.startswith("targeted project focus for "):
                remove_paragraph(paragraph)
                remove_next_project_lines = 0
                continue
            if text.startswith("targeted project evidence for "):
                remove_paragraph(paragraph)
                remove_next_project_lines = 3
                continue
            if remove_next_project_lines > 0:
                if ":" in paragraph.text and norm(paragraph.text) not in {"experience", "technical skills", "skills", "projects", "ai projects", "education"}:
                    remove_paragraph(paragraph)
                    remove_next_project_lines -= 1
                    continue
                remove_next_project_lines = 0

    def heading_index(items: list, names: set[str]) -> int | None:
        for index, paragraph in enumerate(items):
            if norm(paragraph.text) in names:
                return index
        return None

    try:
        doc = Document(str(template_path))
    except Exception as exc:
        logger.warning("Could not open DOCX resume template %s: %s", template_path, exc)
        return False

    clean_previous_targeted_edits(doc)
    doc.core_properties.title = f"{job_title} at {company}"
    profile_text = " ".join(section_lines("Professional Summary")[:2]).strip()
    if not profile_text:
        profile_text = f"Profile updated for {job_title} at {company} using verified resume evidence."

    skills_text = ", ".join(emphasized_skills[:10])
    _ = selected_projects  # Project evidence is tracked in metadata only for Word templates.

    items = all_paragraphs(doc)
    profile_idx = heading_index(items, {"profile", "professional summary", "summary"})
    if profile_idx is not None:
        for paragraph in items[profile_idx + 1:]:
            if norm(paragraph.text) in {"experience", "technical skills", "skills", "projects", "education"}:
                break
            if paragraph.text.strip():
                replace_text(paragraph, profile_text)
                break

    items = all_paragraphs(doc)
    skills_idx = heading_index(items, {"technical skills", "skills"})
    if skills_idx is not None and skills_text:
        insert_after(items[skills_idx], f"Targeted Focus for {job_title}: {skills_text}")

    try:
        doc.save(str(path))
        return True
    except Exception as exc:
        logger.warning("Could not save tailored DOCX %s: %s", path, exc)
        return False


# ---------------------------------------------------------------------------
# Raw XML DOCX writer (no dependencies)
# ---------------------------------------------------------------------------

def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    """Lightweight DOCX writer with no external dependencies."""
    path.parent.mkdir(parents=True, exist_ok=True)
    # Strip heading markers for plain XML
    clean = [p.replace("__HEADING__", "") for p in paragraphs]
    body = "".join(
        f"<w:p><w:r><w:t>{html.escape(paragraph)}</w:t></w:r></w:p>"
        for paragraph in [title, *clean]
        if paragraph
    )
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{body}<w:sectPr /></w:body>
</w:document>"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/document.xml", document_xml)


# ---------------------------------------------------------------------------
# PDF writer (no dependencies)
# ---------------------------------------------------------------------------

def write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    page_width = 612
    page_height = 792
    margin_x = 54
    top_y = 744
    bottom_y = 54
    leading = 14
    content_width = page_width - (margin_x * 2)
    pages: list[list[str]] = [[]]
    y = top_y

    def new_page() -> None:
        nonlocal y
        pages.append([])
        y = top_y

    def ensure_space(height: int = leading) -> None:
        if y - height < bottom_y:
            new_page()

    def escape_text(value: str) -> str:
        clean = value.replace("•", "-")
        clean = clean.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        return clean

    def draw_text(value: str, *, x: int = margin_x, font: str = "F1", size: float = 10.5) -> None:
        pages[-1].append(f"BT /{font} {size:g} Tf {x} {y} Td ({escape_text(value)}) Tj ET")

    def draw_rule() -> None:
        rule_y = y - 4
        pages[-1].append(f"{margin_x} {rule_y} m {page_width - margin_x} {rule_y} l S")

    def wrap(value: str, *, indent: int = 0, size: float = 10.5) -> list[str]:
        chars = max(32, int((content_width - indent) / (size * 0.48)))
        return textwrap.wrap(value.strip(), width=chars, break_long_words=False) or [""]

    ensure_space(30)
    title_lines = wrap(title, size=15)
    for line in title_lines[:2]:
        draw_text(line, x=margin_x, font="F2", size=15)
        y -= 18
    y -= 6

    for raw in paragraphs:
        line = raw.strip()
        if not line:
            y -= 5
            continue

        if line.startswith("__HEADING__"):
            ensure_space(28)
            heading = line[len("__HEADING__"):].strip().upper()
            y -= 5
            draw_text(heading, font="F2", size=10.8)
            draw_rule()
            y -= 16
            continue

        bullet = line.startswith(("•", "-"))
        text = line.lstrip("•-").strip() if bullet else line
        indent = 18 if bullet else 0
        for index, wrapped in enumerate(wrap(text, indent=indent)):
            ensure_space(leading)
            prefix = "- " if bullet and index == 0 else "  " if bullet else ""
            draw_text(f"{prefix}{wrapped}", x=margin_x + indent, size=10.3)
            y -= leading
        if not bullet:
            y -= 2

    page_objects: list[tuple[int, int, bytes]] = []
    next_id = 5
    for page_ops in pages:
        stream = "\n".join(page_ops).encode("latin-1", errors="replace")
        page_objects.append((next_id, next_id + 1, stream))
        next_id += 2

    kids = " ".join(f"{page_id} 0 R" for page_id, _, _ in page_objects)
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_objects)} >>".encode(),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>",
    ]
    for page_id, content_id, stream in page_objects:
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 3 0 R /F2 4 0 R >> >> /Contents {content_id} 0 R >>"
            ).encode()
        )
        objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")

    output = [b"%PDF-1.4\n"]
    offsets: list[int] = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(sum(len(part) for part in output))
        output.append(f"{i} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref_offset = sum(len(part) for part in output)
    output.append(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets:
        output.append(f"{offset:010d} 00000 n \n".encode())
    output.append(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    )
    path.write_bytes(b"".join(output))


def compile_latex_to_pdf(tex_path: Path, pdf_path: Path) -> bool:
    """Compile a LaTeX resume to PDF when a local TeX engine is available."""
    engines: list[list[str]] = []
    if shutil.which("latexmk"):
        engines.append(["latexmk", "-pdf", "-interaction=nonstopmode", "-halt-on-error"])
    for engine in ["xelatex", "pdflatex", "lualatex"]:
        executable = shutil.which(engine)
        if executable:
            engines.append([executable, "-interaction=nonstopmode", "-halt-on-error"])

    if not engines:
        return False

    with tempfile.TemporaryDirectory(prefix="seekapply_latex_") as tmp:
        tmp_path = Path(tmp)
        for engine in engines:
            command = [*engine, "-output-directory", str(tmp_path), tex_path.name]
            try:
                result = subprocess.run(
                    command,
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=45,
                    cwd=str(tex_path.parent),
                )
            except Exception as exc:
                logger.warning("LaTeX compilation could not start with %s: %s", engine[0], exc)
                continue
            generated = tmp_path / f"{tex_path.stem}.pdf"
            if result.returncode == 0 and generated.exists():
                pdf_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copyfile(generated, pdf_path)
                return True
            logger.warning("LaTeX compilation failed with %s: %s", engine[0], result.stderr[-1000:])
    return False


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert DOCX to PDF with a real document renderer when available.

    LibreOffice is preferred for headless servers. On macOS development
    machines, Microsoft Word can be enabled explicitly with
    SEEKAPPLY_ENABLE_MS_WORD_PDF=1. It is disabled by default because Word's
    AppleScript PDF export is slow and can hang or reject commands on some
    installations.
    """
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not docx_path.exists():
        return False

    if executable:
        with tempfile.TemporaryDirectory(prefix="seekapply_docx_pdf_") as tmp:
            tmp_path = Path(tmp)
            try:
                result = subprocess.run(
                    [
                        executable,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(tmp_path),
                        str(docx_path),
                    ],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=45,
                )
            except Exception as exc:
                logger.warning("DOCX to PDF conversion could not start with %s: %s", executable, exc)
            else:
                generated = tmp_path / f"{docx_path.stem}.pdf"
                if result.returncode == 0 and generated.exists():
                    pdf_path.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(generated, pdf_path)
                    return True
                logger.warning("DOCX to PDF conversion failed for %s: %s", docx_path, result.stderr or result.stdout)

    if _microsoft_word_available() and _convert_docx_to_pdf_with_word(docx_path, pdf_path):
        return True
    return False


def _convert_docx_to_pdf_with_word(docx_path: Path, pdf_path: Path) -> bool:
    script = """
on run argv
  set inputPath to item 1 of argv
  set outputPath to item 2 of argv
  set inputFile to POSIX file inputPath
  tell application "Microsoft Word"
    launch
    set visible to false
    set theDoc to open inputFile
    save as theDoc file name outputPath file format format PDF
    close theDoc saving no
  end tell
end run
"""
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        result = subprocess.run(
            ["osascript", "-e", script, str(docx_path), str(pdf_path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=25,
        )
    except Exception as exc:
        logger.warning("Microsoft Word DOCX to PDF conversion could not start: %s", exc)
        return False
    if result.returncode == 0 and pdf_path.exists():
        return True
    logger.warning("Microsoft Word DOCX to PDF conversion failed for %s: %s", docx_path, result.stderr or result.stdout)
    return False


# ---------------------------------------------------------------------------
# LaTeX writer (uses user's Overleaf template)
# ---------------------------------------------------------------------------

def write_latex(
    path: Path,
    template_source: str | None,
    user_name: str,
    skills: list[str],
    experience: list[dict],
    projects: list[dict],
    job_title: str,
    company: str,
    paragraphs: list[str] | None = None,
) -> None:
    """Generate a tailored LaTeX resume.

    If the user uploaded a LaTeX/Overleaf template, keep that template and
    replace common resume sections. Otherwise, create a compact ATS-friendly
    LaTeX resume so the LaTeX download button always has a real file.
    """
    path.parent.mkdir(parents=True, exist_ok=True)

    sectioned = _paragraph_sections(paragraphs or [])
    has_user_template = bool(template_source)
    tex = template_source or _default_latex_template(user_name, job_title, company)

    summary_block = _latex_summary_block(
        sectioned.get("professional summary") or sectioned.get("summary") or sectioned.get("preamble") or []
    )

    # Uploaded LaTeX resumes are treated as the formatting source of truth.
    # Tune only the two highest-signal sections so downloads keep the user's
    # original layout, spacing, projects, education, and experience.
    if has_user_template:
        if summary_block:
            tex = _replace_user_template_summary(tex, summary_block)
        paragraph_skills = _skills_from_paragraphs(sectioned)
        tex = _tune_user_template_skills(tex, skills or paragraph_skills)
        tex = _tune_user_template_projects(tex, projects)
        path.write_text(tex, encoding="utf-8")
        return

    if summary_block:
        tex = _replace_section(tex, "Professional Summary", summary_block)
        tex = _replace_section(tex, "Summary", summary_block)
        tex = _replace_section(tex, "Profile", summary_block)

    # ---- Replace Technical Skills section ----
    paragraph_skills = _skills_from_paragraphs(sectioned)
    skills_block = _latex_skills_block(skills or paragraph_skills)
    tex = _replace_section(tex, "Technical Skills", skills_block)
    tex = _replace_section(tex, "Skills", skills_block)

    # ---- Replace Experience section for generated fallback templates ----
    if experience:
        exp_block = _latex_plain_experience_block(experience)
    else:
        exp_block = _latex_list_block(sectioned.get("experience", []))
    tex = _replace_section(tex, "Experience", exp_block)

    # ---- Replace Projects section for generated fallback templates ----
    if projects:
        proj_block = _latex_plain_projects_block(projects)
    else:
        proj_block = _latex_list_block(sectioned.get("projects", []))
    tex = _replace_section(tex, "Projects", proj_block)
    tex = _replace_section(tex, "AI Projects", proj_block)

    education_block = _latex_list_block(sectioned.get("education", []))
    if education_block:
        tex = _replace_section(tex, "Education", education_block)

    path.write_text(tex, encoding="utf-8")


def _default_latex_template(user_name: str, job_title: str, company: str) -> str:
    return rf"""\documentclass[10pt,letterpaper]{{article}}
\usepackage[margin=0.55in]{{geometry}}
\usepackage[hidelinks]{{hyperref}}
\usepackage{{enumitem}}
\setlength{{\parindent}}{{0pt}}
\setlist[itemize]{{leftmargin=0.16in,itemsep=1pt,topsep=2pt}}
\begin{{document}}
\begin{{center}}
    {{\Large \textbf{{{_latex_escape(user_name)}}}}}\\
    \small Tailored for {_latex_escape(job_title)} at {_latex_escape(company)}
\end{{center}}
\vspace{{-8pt}}
\section*{{Professional Summary}}
\section*{{Technical Skills}}
\section*{{Experience}}
\section*{{Projects}}
\section*{{Education}}
\end{{document}}
"""


def _replace_section(tex: str, section_name: str, new_content: str) -> str:
    """Replace a \\section{Name} or \\section*{Name} block with new content."""
    import re
    pattern = (
        r"(\\section\*?\{" + re.escape(section_name) + r"\})"
        r"(.*?)"
        r"(?=\\section\*?\{|\\end\{document\})"
    )
    match = re.search(pattern, tex, re.DOTALL)
    if match:
        replacement = match.group(1) + "\n" + new_content + "\n"
        tex = tex[:match.start()] + replacement + tex[match.end():]
    return tex


def _replace_user_template_summary(tex: str, new_content: str) -> str:
    for section_name in ["Professional Summary", "Summary", "Profile"]:
        updated = _replace_section_preserving_wrappers(tex, section_name, new_content)
        if updated != tex:
            return updated
    return tex


def _replace_section_preserving_wrappers(tex: str, section_name: str, new_content: str) -> str:
    """Replace one section body while keeping common resume spacing wrappers."""
    import re

    pattern = (
        r"(\\section\*?\{" + re.escape(section_name) + r"\})"
        r"(.*?)"
        r"(?=\\section\*?\{|\\end\{document\})"
    )
    match = re.search(pattern, tex, re.DOTALL)
    if not match:
        return tex

    block = match.group(2)

    small_match = re.search(r"\\small\s*\{", block)
    if small_match:
        open_brace = block.find("{", small_match.start())
        close_brace = _find_latex_group_end(block, open_brace)
    else:
        open_brace = -1
        close_brace = -1

    if small_match and open_brace >= 0 and close_brace >= 0:
        replacement_block = (
            block[: open_brace + 1]
            + "\n"
            + new_content
            + "\n"
            + block[close_brace:]
        )
    else:
        leading_match = re.match(r"(\s*(?:\\vspace\{[^}]+\}\s*)*)", block)
        trailing_match = re.search(r"((?:\s*\\vspace\{[^}]+\})\s*)$", block)
        leading = leading_match.group(1) if leading_match else "\n"
        trailing = trailing_match.group(1) if trailing_match else "\n"
        replacement_block = f"{leading}{new_content}{trailing}"

    replacement = match.group(1) + replacement_block
    return tex[: match.start()] + replacement + tex[match.end():]


def _find_latex_group_end(text: str, open_brace_index: int) -> int:
    """Return the matching closing brace for a LaTeX group."""
    if open_brace_index < 0 or open_brace_index >= len(text) or text[open_brace_index] != "{":
        return -1

    depth = 0
    index = open_brace_index
    while index < len(text):
        char = text[index]
        if char == "\\":
            index += 2
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index
        index += 1
    return -1


def _tune_user_template_skills(tex: str, skills: list[str]) -> str:
    focus_skills = [skill for skill in skills if skill][:10]
    if not focus_skills:
        return tex
    focus = ", ".join(_latex_escape(skill) for skill in focus_skills)
    focus_line = rf"\textbf{{Targeted Focus:}} {focus} \\"

    for section_name in ["Technical Skills", "Skills"]:
        updated = _replace_or_insert_focus_line(tex, section_name, focus_line)
        if updated != tex:
            return updated
    return tex


def _replace_or_insert_focus_line(tex: str, section_name: str, focus_line: str) -> str:
    import re

    pattern = (
        r"(\\section\*?\{" + re.escape(section_name) + r"\})"
        r"(.*?)"
        r"(?=\\section\*?\{|\\end\{document\})"
    )
    match = re.search(pattern, tex, re.DOTALL)
    if not match:
        return tex

    block = match.group(2)
    if r"\textbf{Targeted Focus:}" in block:
        block = re.sub(
            r"\\textbf\{Targeted Focus:\}.*?(?:\\\\)?",
            focus_line,
            block,
            count=1,
            flags=re.DOTALL,
        )
    elif r"\small{\item{" in block:
        block = block.replace(r"\small{\item{", r"\small{\item{" + "\n    " + focus_line + "\n    ", 1)
    elif r"\item{" in block:
        block = block.replace(r"\item{", r"\item{" + "\n    " + focus_line + "\n    ", 1)
    elif r"\item " in block:
        block = re.sub(r"(\\item\s+)", r"\1" + focus_line + "\n    ", block, count=1)
    else:
        block = "\n" + _latex_focus_itemize(focus_line) + block

    replacement = match.group(1) + block
    return tex[: match.start()] + replacement + tex[match.end():]


def _tune_user_template_projects(tex: str, projects: list[dict]) -> str:
    selected = [project for project in projects if project.get("name")]
    if not selected:
        return tex
    for section_name in ["AI Projects", "Projects"]:
        current = _section_body(tex, section_name)
        if current is None:
            continue
        additions = _new_projects_only(current, selected)
        if not additions:
            return tex
        if r"\resumeProjectHeading" in current:
            block = _append_project_macro_entries(current, additions)
        elif r"\begin{itemize}" in current:
            block = _append_plain_project_items(current, additions)
        else:
            continue
        updated = _replace_section(tex, section_name, block)
        if updated != tex:
            return updated
    return tex


def _section_body(tex: str, section_name: str) -> str | None:
    import re

    pattern = (
        r"\\section\*?\{" + re.escape(section_name) + r"\}"
        r"(.*?)"
        r"(?=\\section\*?\{|\\end\{document\})"
    )
    match = re.search(pattern, tex, re.DOTALL)
    return match.group(1) if match else None


def _new_projects_only(current_section: str, projects: list[dict]) -> list[dict]:
    current = _normalize_latex_text(current_section)
    selected: list[dict] = []
    for project in projects:
        name = _normalize_latex_text(str(project.get("name") or ""))
        url = _normalize_latex_text(str(project.get("url") or project.get("repo_url") or ""))
        if not name:
            continue
        if name in current or (url and url in current):
            continue
        selected.append(project)
        if len(selected) >= 3:
            break
    return selected


def _normalize_latex_text(value: str) -> str:
    import re

    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _append_project_macro_entries(current_section: str, projects: list[dict]) -> str:
    entries = _latex_project_macro_entries(projects)
    if r"\resumeSubHeadingListEnd" in current_section:
        return current_section.replace(r"\resumeSubHeadingListEnd", entries + "\n    " + r"\resumeSubHeadingListEnd", 1)
    return current_section.rstrip() + "\n" + entries + "\n"


def _append_plain_project_items(current_section: str, projects: list[dict]) -> str:
    entries = "\n".join(_latex_plain_project_item(project) for project in projects)
    if r"\end{itemize}" in current_section:
        return current_section.replace(r"\end{itemize}", entries + "\n" + r"\end{itemize}", 1)
    return current_section.rstrip() + "\n" + entries + "\n"


def _latex_project_macro_entries(projects: list[dict]) -> str:
    lines: list[str] = []
    for project in projects:
        name = _latex_escape(str(project.get("name", "Project"))[:80])
        skills = [str(skill) for skill in project.get("skills", []) if str(skill).strip()]
        skill_line = ", ".join(_latex_escape(skill) for skill in skills[:9]) or "Verified project evidence"
        url = str(project.get("url") or project.get("repo_url") or "").strip()
        link = rf"{{\href{{{url}}}{{\faGithub}}}}" if url.startswith(("http://", "https://")) else "{}"
        bullets = [str(item).strip() for item in project.get("bullets", []) if str(item).strip()]
        summary = str(project.get("summary") or project.get("description") or "").strip()
        if summary:
            bullets.insert(0, summary)
        bullets = bullets[:3] or ["Project evidence retained from the uploaded resume or GitHub project notes."]
        lines.extend(
            [
                r"    \resumeProjectHeading",
                rf"      {{\textbf{{{name}}} $|$ \emph{{{skill_line}}}}}",
                rf"      {link}",
                r"    \resumeItemListStart",
            ]
        )
        for bullet in bullets:
            lines.append(r"      \resumeItem{" + _latex_escape(bullet) + "}")
        lines.extend([r"    \resumeItemListEnd", r"    \vspace{-5pt}", ""])
    return "\n".join(lines)


def _latex_plain_project_item(project: dict) -> str:
    name = _latex_escape(str(project.get("name", "Project"))[:80])
    summary = _latex_escape(str(project.get("summary") or project.get("description") or "Verified GitHub project evidence."))
    skills = ", ".join(_latex_escape(str(skill)) for skill in (project.get("skills") or [])[:8])
    details = summary + (f" ({skills})" if skills else "")
    return r"    \item \textbf{" + name + r"}: " + details


def _latex_focus_itemize(focus_line: str) -> str:
    return (
        r"\begin{itemize}[leftmargin=0.1in, label={}]" "\n"
        r"    \item " + focus_line + "\n"
        r"\end{itemize}" "\n"
    )


def _paragraph_sections(paragraphs: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = "preamble"
    for raw in paragraphs:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("__HEADING__"):
            current = line[len("__HEADING__"):].strip().lower()
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return sections


def _skills_from_paragraphs(sections: dict[str, list[str]]) -> list[str]:
    for key in ["skills", "relevant skills", "technical skills"]:
        lines = sections.get(key, [])
        if lines:
            return [item.strip() for item in ",".join(lines).split(",") if item.strip()]
    return []


def _latex_summary_block(lines: list[str]) -> str:
    summary = " ".join(line.strip("•- ") for line in lines if line.strip())
    return _latex_escape(summary) if summary else ""


def _latex_escape(text: str) -> str:
    """Escape special LaTeX characters."""
    normalized = []
    for char in text.replace("\t", " ").replace("•", "-"):
        if ord(char) < 128:
            normalized.append(char)
        elif char in {"–", "—", "−"}:
            normalized.append("-")
        elif char in {"“", "”"}:
            normalized.append('"')
        elif char in {"‘", "’"}:
            normalized.append("'")
    text = "".join(normalized)
    replacements = {
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for char, escaped in replacements.items():
        text = text.replace(char, escaped)
    return text


def _latex_skills_block(skills: list[str]) -> str:
    """Generate the Technical Skills section content."""
    skills_line = ", ".join(_latex_escape(s) for s in skills)
    if not skills_line:
        skills_line = "Skills available in the uploaded resume"
    return (
        r"\begin{itemize}[leftmargin=0.1in, label={}]" "\n"
        r"    \item \textbf{Technical Skills:} " + skills_line + "\n"
        r"\end{itemize}" "\n"
        r"\vspace{-8pt}"
    )


def _latex_list_block(lines: list[str]) -> str:
    cleaned = [line.strip() for line in lines if line.strip()]
    if not cleaned:
        return ""
    output = [r"\begin{itemize}"]
    for line in cleaned[:12]:
        output.append(r"    \item " + _latex_escape(line.lstrip("•- ")))
    output.append(r"\end{itemize}")
    output.append(r"\vspace{-8pt}")
    return "\n".join(output)


def _latex_plain_experience_block(experience: list[dict]) -> str:
    output: list[str] = []
    for exp in experience[:4]:
        company = _latex_escape(str(exp.get("company", "")))
        role = _latex_escape(str(exp.get("role", "")))
        duration = _latex_escape(str(exp.get("duration", "")))
        location = _latex_escape(str(exp.get("location", "")))
        output.append(r"\textbf{" + role + r"} \hfill " + duration + r"\\")
        output.append(company + (f" -- {location}" if location else "") + "\n")
        bullets = [str(bullet) for bullet in exp.get("bullets", [])[:5]]
        if bullets:
            output.append(r"\begin{itemize}")
            for bullet in bullets:
                output.append(r"    \item " + _latex_escape(bullet))
            output.append(r"\end{itemize}")
    return "\n".join(output)


def _latex_plain_projects_block(projects: list[dict]) -> str:
    output = [r"\begin{itemize}"]
    for proj in projects[:5]:
        name = _latex_escape(str(proj.get("name", "Project")))
        summary = _latex_escape(str(proj.get("summary", "")))
        skills = ", ".join(_latex_escape(str(skill)) for skill in proj.get("skills", []))
        details = summary
        if skills:
            details = f"{details} ({skills})" if details else skills
        output.append(r"    \item \textbf{" + name + r"}: " + details)
    output.append(r"\end{itemize}")
    output.append(r"\vspace{-8pt}")
    return "\n".join(output)


def _latex_experience_block(experience: list[dict]) -> str:
    """Generate the Experience section content."""
    lines = [r"  \resumeSubHeadingListStart"]
    for exp in experience[:4]:
        company = _latex_escape(exp.get("company", ""))
        role = _latex_escape(exp.get("role", ""))
        duration = _latex_escape(exp.get("duration", ""))
        location = _latex_escape(exp.get("location", ""))
        lines.append(f"    \\resumeSubheading")
        lines.append(f"      {{{company}}}{{{duration}}}")
        lines.append(f"      {{{role}}}{{{location}}}")
        lines.append(r"      \resumeItemListStart")
        for bullet in exp.get("bullets", [])[:5]:
            lines.append(f"        \\resumeItem {{{_latex_escape(bullet)}}}")
        lines.append(r"    \resumeItemListEnd")
        lines.append("")
    lines.append(r"  \resumeSubHeadingListEnd")
    lines.append(r"\vspace{-16pt}")
    return "\n".join(lines)


def _latex_projects_block(projects: list[dict]) -> str:
    """Generate the Projects section content."""
    lines = [r"    \vspace{-6pt}", r"    \resumeSubHeadingListStart"]
    for proj in projects[:5]:
        name = _latex_escape(proj.get("name", ""))
        url = proj.get("url", "")
        gh_link = ""
        if url:
            gh_link = r"{\href{" + url + r"}{\faGithub}}"
        lines.append(f"    \\resumeProjectHeading")
        lines.append(f"          {{\\textbf{{{name}   }}{gh_link}}}{{}}")
        lines.append(r"          \resumeItemListStart")
        summary = proj.get("summary", "")
        if summary:
            for part in summary.split(";"):
                part = part.strip()
                if part:
                    lines.append(f"            \\resumeItem{{{_latex_escape(part)}}}")
        skills_list = proj.get("skills", [])
        if skills_list:
            skills_str = ", ".join(_latex_escape(s) for s in skills_list)
            lines.append(f"            \\resumeItem{{Technologies Used: {skills_str}}}")
        lines.append(r"          \resumeItemListEnd")
        lines.append(r"          \vspace{-14pt}")
        lines.append("")
    lines.append(r"    \resumeSubHeadingListEnd")
    lines.append(r"\vspace{3pt}")
    return "\n".join(lines)
